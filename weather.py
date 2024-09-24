
import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime, timedelta
import plotly.graph_objs as go
import matplotlib.pyplot as plt
import logging
import plotly.figure_factory as ff
import altair as alt
import geopandas as gpd
import folium
from streamlit_folium import folium_static
from folium import Popup, GeoJson
import io
import zipfile
from docx import Document
import base64

logging.basicConfig(level=logging.INFO)
st.set_page_config(layout="wide")
PASSWORD = "Abinbev@123"

def password_protection():
    def check_password():
        if st.session_state["password"] == PASSWORD:
            st.session_state["password_correct"] = True
            st.session_state["password_attempts"] = 0
        else:
            st.session_state["password_correct"] = False
            st.session_state["password_attempts"] += 1
    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False
        st.session_state["password_attempts"] = 0
    if not st.session_state["password_correct"]:
        if st.session_state["password_attempts"] >= 3:
            st.error("Too many incorrect attempts. Please refresh the page to try again.")
        else:
            st.text_input("Enter password:", type="password", on_change=check_password, key="password")
            if st.session_state["password_correct"]:
                st.success("Password correct! Access granted.")
            elif st.session_state["password_attempts"] > 0:
                st.error("Password incorrect. Please try again.")
        st.stop()
password_protection()

st.markdown(
    """
    <style>
    body {
        background-color: #ffffff;
        color: #333;
    }
    .stApp {
        background-color: #ffffff;
    }
    .box-container {
        display: flex;
        justify-content: space-around;
        flex-wrap: wrap;
    }
    .box {
        background-color: #ffffff;
        border: 2px solid #e0e0e0;
        border-radius: 10px;
        padding: 20px;
        margin: 10px;
        text-align: center;
        box-shadow: 2px 2px 5px rgba(0, 0, 0, 0.1);
        width: 150px;
        height: 180px;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
    }
    .box h2 {
        color: #333;
        margin-bottom: 10px;
        font-size: 1em;
    }
    .box p {
        font-size: 1.5em;
        color: #555;
        margin: 0;
    }
    </style>
    """,
    unsafe_allow_html=True
)

def create_map(gdf, zoom_level):
    if gdf is None:
        return None

    center = [gdf.geometry.centroid.y.mean(), gdf.geometry.centroid.x.mean()]
    m = folium.Map(location=center, zoom_start=zoom_level, control_scale=True)
    geojson_data = gdf.__geo_interface__

    # Update colors to reflect mature stages instead of severity
    mature_seed_colors = {
        ("Not Harvesting", "1207"): "#ff4c4c",    # Red for Late Stage and 1207
        ("Not Harvesting", "8214"): "#e53d3d",    # Slightly darker red for Late Stage and 8214
        ("Not Harvesting", "R&D Plot"): "#c62828", # Even darker red for Late Stage and R&D Plot

        ("Harvesting", "1207"): "#fbc02d",     # Yellow for Mid Stage and 1207
        ("Harvesting", "8214"): "#f9a825",     # Slightly darker yellow for Mid Stage and 8214
        ("Harvesting", "R&D Plot"): "#f57f17", # Even darker yellow for Mid Stage and R&D Plot

        ("Threshing", "1207"): "#66bb6a",   # Green for Early Stage and 1207
        ("Threshing", "8214"): "#43a047",   # Slightly darker green for Early Stage and 8214
        ("Threshing", "R&D Plot"): "#388e3c" # Even darker green for Early Stage and R&D Plot
    }

    for idx, feature in enumerate(geojson_data['features']):
        # Use Mature Stage instead of Severity Level
        mature_stage = feature['properties'].get('Mature Stage', '')
        seed_variety = feature['properties'].get('Seed-Varity', 'N/A')
        
        color = mature_seed_colors.get((mature_stage, seed_variety), "gray")
        
        simple_fields = {
            "Name": feature['properties'].get('Name', 'N/A'),
            "Area (Bigha)": feature['properties'].get('Area (Bigha)', 'N/A'),
            "Sowing dates": feature['properties'].get('Sowing dates', 'N/A'),
            "Harvesting Date": feature['properties'].get('Harvesting Date', 'N/A'),
            "Yield (kg/bigha)": feature['properties'].get('Yield (kg/bigha)', 'N/A'),
            "Seed Variety": seed_variety
        }
        
        all_fields = feature['properties']
        keys = list(all_fields.keys())
        values = list(all_fields.values())
        split_index = len(keys) // 2
        column1 = ''.join(f'<b>{keys[i]}</b>: {values[i]}<br>' for i in range(split_index))
        column2 = ''.join(f'<b>{keys[i]}</b>: {values[i]}<br>' for i in range(split_index, len(keys)))
        simple_tooltip_html = ''.join(f'<b>{key}</b>: {value}<br>' for key, value in simple_fields.items())
        detailed_tooltip_html = f'''
        <div style="display: flex;">
            <div style="width: 50%; padding-right: 15px;">{column1}</div>
            <div style="width: 50%; padding-left: 15px;">{column2}</div>
        </div>
        '''
        
        doc_buffer = create_docx_for_plot(feature['properties'])
        doc_base64 = base64.b64encode(doc_buffer.getvalue()).decode()
        plot_name = feature['properties'].get('Name', f'plot_{idx + 1}').replace(' ', '_')
        download_link = f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{doc_base64}'
        tooltip_html = f'''
        <div style="display: flex; flex-wrap: wrap; width: 600px; font-size: 14px; background-color: #fff; border-radius: 8px; padding: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            <div id="simpleView" style="display: block;">{simple_tooltip_html}<a href="#" onclick="document.getElementById('simpleView').style.display='none';document.getElementById('detailedView').style.display='block';return false;">Show Details</a></div>
            <div id="detailedView" style="display: none;">{detailed_tooltip_html}</div>
            <div><a href="{download_link}" download="{plot_name}.docx">Download DOCX</a></div>
        </div>
        '''
        
        folium.GeoJson(
            feature,
            style_function=lambda x, color=color: {'fillColor': color, 'color': color, 'weight': 2, 'fillOpacity': 0.6},
            popup=Popup(tooltip_html, max_width=600)
        ).add_to(m)
    
    folium.LayerControl().add_to(m)
    return m
    
def load_geojson(file_path):
    try:
        gdf = gpd.read_file(file_path)
        if gdf.empty:
            st.error("GeoDataFrame is empty. Please check the GeoJSON file.")
        for column in gdf.select_dtypes(include=['datetime', 'timedelta']).columns:
            gdf[column] = gdf[column].dt.strftime('%Y-%m-%d')
        return gdf
    except Exception as e:
        st.error(f"Error loading GeoJSON file: {e}")
        return None

def create_bar_plots(df):
    for column in df.columns:
        if column not in ['Tillage Operation', 'Farm Name']:
            st.subheader(f"{column}")
            total_count = df.shape[0]
            done_count = df[df[column].isin(['Done', 'Done early', 'Done on time'])].shape[0]
            df_plot = pd.DataFrame({
                'Status': ['Total', 'Done and Pop Followed'],
                'Count': [total_count, done_count]
            })
            colors = ['blue', 'green']
            fig = px.bar(df_plot, x='Count', y='Status', orientation='h', text='Count', width=800, height=200, color='Status', color_discrete_sequence=colors)
            fig.update_traces(texttemplate='%{text}', textposition='outside')
            fig.update_layout(
                xaxis_title="Count",
                yaxis_title="Status",
                title=f"{column}",
                showlegend=False
            )
            st.plotly_chart(fig)
            
def get_color(severity):
    if severity == "Low":
        return "green"
    elif severity == "Medium":
        return "yellow"
    elif severity == "High":
        return "red"
    else:
        return "gray"

def create_docx_for_plot(plot_data):
    doc = Document()
    doc.add_heading(f"Plot Details - {plot_data.get('Name', 'Unnamed Plot')}", 0)
    for key, value in plot_data.items():
        if key != "geometry":
            doc.add_paragraph(f"{key}: {value}")
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def create_zip_file(gdf):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for idx, row in gdf.iterrows():
            plot_data = row.to_dict()
            doc_buffer = create_docx_for_plot(plot_data)
            plot_name = plot_data.get('Name', f'plot_{idx + 1}').replace(' ', '_')
            zipf.writestr(f"{plot_name}.docx", doc_buffer.getvalue())
    zip_buffer.seek(0)
    return zip_buffer
            
def plot_severity_counts(df, sort_by='specific_order'):
    df['Date'] = pd.to_datetime(df['Date'], dayfirst=True)
    latest_records = df.loc[df.groupby('farmName')['Date'].idxmax()]
    severity_counts = latest_records['Severity'].value_counts().reset_index()
    severity_counts.columns = ['Severity', 'Count']
    color_scale = alt.Scale(
        domain=['Low', 'medium', 'high', 'medium-large area affected'],
        range=['green', 'yellow', 'red', 'orange']  # Add orange for the new category
    )
    if sort_by == 'specific_order':
        severity_order = ['Low', 'medium', 'high', 'medium-large area affected']
        x_encoding = alt.X('Severity:O', sort=severity_order)
    elif sort_by == 'count':
        severity_counts = severity_counts.sort_values(by='Count', ascending=False)
        x_encoding = alt.X('Severity:O', sort='-y')
    else:
        raise ValueError("sort_by must be either 'specific_order' or 'count'")
    bar_chart = alt.Chart(severity_counts).mark_bar().encode(
        x=x_encoding,
        y='Count:Q',
        color=alt.Color('Severity:N', scale=color_scale),
        tooltip=['Severity', 'Count']
    ).properties(
        title='Count of Severity Levels in Latest Records for Each Farm'
    )
    text = bar_chart.mark_text(
        align='center',
        baseline='middle',
        dx=0,  # shift text horizontally to center
        dy=-5  # shift text vertically to center
    ).encode(
        text=alt.Text('Count:Q', format='.0f'),  # format to integer
    )
    bar_chart = (bar_chart + text).properties(
        width=alt.Step(60)
    )
    st.altair_chart(bar_chart, use_container_width=True)

def create_activity_progress_plot():
    data = [
        dict(Task='DAP / MOP Fertilizer', Start='2024-05-15', Finish='2024-05-20', Done=59, NotDone=0, Status='done'),
        dict(Task='Sowing', Start='2024-05-15', Finish='2024-05-20', Done=59, NotDone=0, Status='done'),
        dict(Task='Weeding 1', Start='2024-06-05', Finish='2024-06-11', Done=59, NotDone=0, Status='done'),
        dict(Task='Irrigation 1', Start='2024-06-13', Finish='2024-06-15', Done=59, NotDone=0, Status='done'),
        dict(Task='Urea 1', Start='2024-06-15', Finish='2024-06-17', Done=59, NotDone=0, Status='done'),
        dict(Task='Weeding 2', Start='2024-06-28', Finish='2024-07-04', Done=17, NotDone=42, Status='done'),
        dict(Task='Irrigation 2', Start='2024-07-01', Finish='2024-07-05', Done=59, NotDone=0, Status='done'),
        dict(Task='Fungicide Spray-1', Start='2024-07-16', Finish='2024-07-25', Done=13, NotDone=46, Status='done'),
        dict(Task='Herbicide Spray', Start='2024-07-09', Finish='2024-07-12', Done=59, NotDone=0, Status='done'),
        dict(Task='Irrigation 3', Start='2024-07-17', Finish='2024-07-22', Done=59, NotDone=0, Status='done'),
        dict(Task='Urea 2', Start='2024-07-20', Finish='2024-07-24', Done=59, NotDone=0, Status='done'),
        dict(Task='Fungicide Spray-2', Start='2024-08-03', Finish='2024-08-08', Done=59, NotDone=0, Status='done'),
        dict(Task='Pesticide Spray', Start='2024-08-09', Finish='2024-08-10', Done=8, NotDone=51, Status='done'),
        dict(Task='Fungicide Spray-3', Start='2024-09-06', Finish='2024-09-07', Done=1, NotDone=58, Status='done'),
        dict(Task='Harvesting', Start='2024-09-08', Finish='2024-09-24', Done=47, NotDone=12, Status='ongoing'),
        dict(Task='Threshing', Start='2024-09-16', Finish='2024-09-24', Done=15, NotDone=44, Status='ongoing')
    ]
    data_combined = []
    for item in data:
        if item['NotDone'] == 0:
            color = 'green'
        elif item['Status'] == 'ongoing':
            color = 'red'
        else:
            color = 'orange'      
        hover_text = (f"Task: {item['Task']}<br>"
                      f"Start: {item['Start']}<br>"
                      f"Finish: {item['Finish']}<br>"
                      f"Done: {item['Done']}<br>"
                      f"NotDone: {item['NotDone']}<br>"
                      f"NotDone Count: {item['NotDone']}")  # Adding NotDone Count        
        data_combined.append(dict(Task=item['Task'], Start=item['Start'], Finish=item['Finish'], Resource=color, Done=item['Done'], NotDone=item['NotDone'], hovertext=hover_text))
    fig = ff.create_gantt(data_combined, index_col='Resource', group_tasks=True, showgrid_x=True, showgrid_y=True, colors={'green': 'rgb(0, 255, 0)', 'orange': 'rgb(255, 165, 0)', 'red': 'rgb(255, 0, 0)'})
    fig.update_layout(
        xaxis_title="Date",
        yaxis_title="Activity"
    )
    st.plotly_chart(fig)

def severity_dot_plot(data):
    data['Date'] = pd.to_datetime(data['Date'], format='%d/%m/%Y')
    color_map = {
        'Low': 'green',
        'medium': 'yellow',
        'high': 'red',
        'medium-large area affected': 'orange'  # New category
    }
    medium_high_large = data[data['Severity'].isin(['medium', 'high', 'medium-large area affected'])]
    low_all = data[data['Severity'] == 'Low']
    filtered_data = pd.concat([medium_high_large, low_all]).sort_values('Date')
    filtered_data['Color'] = filtered_data['Severity'].map(color_map)
    filtered_data['Farm Name HTML'] = filtered_data.apply(
        lambda row: f"<a href='https://riskchart.streamlit.app/?farmName={row['farmName'].replace(' ', '+')}' target='_blank'>{row['farmName']}</a>",
        axis=1
    )
    filtered_data['CustomData'] = filtered_data.apply(
        lambda row: (row['Date'].strftime('%d/%m/%Y'), row['Severity'], row['Note'], row['farmName']),
        axis=1
    )
    fig = px.scatter(
        filtered_data,
        x='Date',  # Plot against Date to show all entries
        y='Farm Name HTML',  # Use HTML links here
        color='Severity',
        color_discrete_map=color_map,
        hover_data={'Date': False, 'Severity': False, 'Note': False, 'Farm Name HTML': False, 'CustomData': False}
    )
    fig.update_traces(
        hovertemplate="<b>Severity:</b> %{customdata[1]}<br><b>Farm Name:</b> %{customdata[3]}<br><b>Date:</b> %{customdata[0]}<br><b>Note:</b> %{customdata[2]}"
    )
    fig.update_layout(
        xaxis_title='Date',
        yaxis_title='Farm Name',
        legend_title='Severity',
        height=1500
    )
    fig.update_traces(marker=dict(size=10), selector=dict(mode='markers'))
    return fig

def create_line_chart(df, time_frame):
    try:
        logging.info("Converting 'Date' column to datetime format.")
        df['Date'] = pd.to_datetime(df['Date'], errors='coerce')       
        if df['Date'].isnull().any():
            logging.warning("There are NaT values in 'Date' column after conversion.")        
        if time_frame == 'Last 2 Days':
            last_days = datetime.now() - timedelta(days=2)
            df_filtered = df[df['Date'] >= last_days]
        elif time_frame == 'Last Week':
            last_days = datetime.now() - timedelta(days=7)
            df_filtered = df[df['Date'] >= last_days]
        elif time_frame == 'Last Month':
            last_days = datetime.now() - timedelta(days=30)
            df_filtered = df[df['Date'] >= last_days]
        elif time_frame == 'Visit Till Date':
            df_filtered = df[df['Date'] <= datetime.now()]
        else:
            raise ValueError("Invalid time_frame value provided.")       
        if df_filtered.empty:
            logging.warning("The filtered dataframe is empty.")        
        logging.info("Grouping by 'FarmName' and 'Date' to count activities.")
        df_counts = df_filtered.groupby(['FarmName', 'Date']).size().reset_index(name='Activity')     
        logging.info("Grouping by 'FarmName' to count unique dates with activities.")
        df_counts = df_counts.groupby('FarmName').size().reset_index(name='Activity')       
        logging.info("Creating the line chart.")
        fig = px.line(df_counts, x='FarmName', y='Activity', title=f'Activities per Plot in {time_frame}')
        fig.update_layout(xaxis_title='FarmName', yaxis_title='Number of Activities')       
        logging.info("Line chart created successfully.")
        return fig    
    except Exception as e:
        logging.error(f"An error occurred: {e}")
        raise

def create_dot_plot(df):
    df = df.astype(str)
    dot_data = []
    for column in df.columns:
        if column != 'Farm Name':
            for idx, value in enumerate(df[column]):
                color = None
                if value.startswith('well and passed'):
                    color = 'green'
                elif value.startswith('current'):
                    color = 'yellow'
                elif not value.startswith(('well and passed', 'current')):
                    color = 'red'
                if color:
                    dot_data.append({'Column': column, 'FarmName': df['Farm Name'][idx], 'Value': value, 'Color': color})
    dot_df = pd.DataFrame(dot_data)
    dot_df = dot_df.sort_values(by='Column')  
    fig = px.scatter(dot_df, x='Column', y='FarmName', color='Color', 
                     color_discrete_map={
                         'green': 'green',
                         'yellow': 'yellow',
                         'red': 'red'
                     },
                     title="Plot wise Growth Stage Summary",
                     labels={'FarmName': 'Farm Name'},
                     hover_data={'Value': True, 'Color': False})
    fig.for_each_trace(lambda t: t.update(name = {'green': 'well and passed', 
                                                  'yellow': 'current stage', 
                                                  'red': 'not in current stage or some Alert'}[t.name]))
    fig.update_layout(height=1500)
    fig.update_traces(marker=dict(size=10), selector=dict(mode='markers'))
    return fig

def create_dot_plot_1(df):
    df = df.astype(str)
    dot_data = []
    columns_to_process = list(df.columns)    
    if 'Farm Name' in columns_to_process:
        columns_to_process.remove('Farm Name')    
    for column in columns_to_process:
        for idx, value in enumerate(df[column]):
            color = None
            value_lower = value.lower().strip()            
            if value_lower == "inprogress":
                color = 'yellow'
            elif column in ['Sowing', 'M0P/DAP', 'UREA 1']:
                try:
                    numeric_value = float(value.replace(',', '').split()[0])
                    if (column == 'Sowing' and 7.5 <= numeric_value <= 8.5) or \
                       (column == 'M0P/DAP' and 9.5 <= numeric_value <= 10.5) or \
                       (column == 'UREA 1' and 6.5 <= numeric_value <= 7.5):
                        color = 'green'
                    else:
                        color = 'red'
                except ValueError:
                    pass
            elif column in ['Weeding 1', 'Irrigation 1', 'Weeding 2', 'Disease_Control_Fungicide_1', 'Irrigation 2', 'Herbiside', 'Irrigation 3', 'UREA 2', 
                            'Disease_Control_Fungicide_2', 'Pesticide_spray', 'Disease_Control_Fungicide_3', 'Harvest 1', 'Threshing'
                        ]:
                if value_lower not in ['nan', '0', 'null', '']:
                    color = 'green'
                else:
                    color = 'red'           
            if color:
                farm_name = df['Farm Name'][idx]
                farm_link = f"<a href='https://farmimage.streamlit.app/?farm_name={farm_name}' target='_blank'>{farm_name}</a>"
                dot_data.append({'Column': column, 'FarmName': farm_link, 'Value': value, 'Color': color})   
    dot_df = pd.DataFrame(dot_data)
    dot_df['Column'] = pd.Categorical(dot_df['Column'], categories=columns_to_process, ordered=True)
    dot_df = dot_df.sort_values(by='Column')  
    fig = px.scatter(dot_df, x='Column', y='FarmName', color='Color',
                     color_discrete_map={'green': 'green', 'red': 'red', 'yellow': 'yellow'},
                     title="Plot wise Activity Summary",
                     category_orders={'Column': columns_to_process},
                     labels={'FarmName': 'Farm Name'},
                     hover_data={'Value': True, 'FarmName': False, 'Color': False})    
    fig.for_each_trace(lambda t: t.update(name={
        'green': 'well and pop Followed', 
        'red': 'pop not followed or Activity not Done',
        'yellow': 'in progress'
    }[t.name]))    
    fig.update_layout(height=1500)
    fig.update_traces(marker=dict(size=10), selector=dict(mode='markers'))   
    return fig

def create_stacked_bar_chart(df):
    df = df.astype(str)
    status_counts = {}
    for column in df.columns:
        if column != 'Farm Name':
            well_and_passed_count = df[df[column].str.startswith('well and passed', na=False)].shape[0]
            current_count = df[df[column].str.startswith('current', na=False)].shape[0]
            other_count = df[~df[column].str.startswith(('well and passed', 'current'), na=False)].shape[0]
            status_counts[column] = {
                'Well and Passed': well_and_passed_count,
                'Current': current_count,
                'Some Alert or not reached to this stage': other_count
            }
    status_df = pd.DataFrame(status_counts).T.reset_index()
    status_df = status_df.melt(id_vars='index', var_name='Status', value_name='Count')
    fig = px.bar(status_df, x='index', y='Count', color='Status', 
                 color_discrete_map={
                     'Well and Passed': 'green', 
                     'Current': 'yellow', 
                     'Some Alert or not reached to this stage': 'red'
                 },
                 title="Growth Tracker Summary",
                 labels={'index': 'Columns'})
    return fig

def kelvin_to_celsius(k):
    return k - 273.15

def plot_weather_trends(weather_df):
    weather_df['dt'] = pd.to_datetime(weather_df['dt'], errors='coerce')
    weather_df['temp_max_C'] = kelvin_to_celsius(weather_df['temp_max'])
    weather_df['temp_min_C'] = kelvin_to_celsius(weather_df['temp_min'])
    weather_df['clouds_normalized'] = weather_df['clouds'] / 100.0
    weather_df['snow_emoji'] = weather_df['snow'].apply(lambda x: '❄️' if x > 0 else '')
    weather_df['rain_emoji'] = weather_df['rain'].apply(lambda x: '🌧️' if x > 0 else '')
    weather_df['clouds_emoji'] = weather_df['clouds_normalized'].apply(lambda x: '☁️' if x > 0.5 else '')
    fig = px.line(weather_df, x='dt', y=['temp_max_C', 'temp_min_C'],
                  title='Temperature Trends',
                  labels={'dt': 'Date', 'value': 'Temperature (°C)'},
                  line_shape='linear')
    fig.add_scatter(x=weather_df['dt'], y=[weather_df['temp_max_C'].max()] * len(weather_df),
                    mode='text', text=weather_df['snow_emoji'], name='Snow', textposition='top center')
    fig.add_scatter(x=weather_df['dt'], y=[weather_df['temp_min_C'].max()] * len(weather_df),
                    mode='text', text=weather_df['rain_emoji'], name='Rain', textposition='top center')
    fig.add_scatter(x=weather_df['dt'], y=[weather_df['temp_min_C'].min()] * len(weather_df),
                    mode='text', text=weather_df['clouds_emoji'], name='Clouds', textposition='top center')
    fig.update_layout(height=700)
    return fig

activity_data_url = "https://raw.githubusercontent.com/sakshamraj4/Abinbav_sustainability/main/activity_avinbev.csv"
activity_df = pd.read_csv(activity_data_url)
new_data_url = "https://raw.githubusercontent.com/sakshamraj4/Abinbav_sustainability/main/data.csv"
new_data_df = pd.read_csv(new_data_url)
growth_tracker_url = "https://raw.githubusercontent.com/sakshamraj4/Abinbav_sustainability/main/Growth_Tracker.csv"
growth_tracker_df = pd.read_csv(growth_tracker_url)
growth_data_csv_url = "https://raw.githubusercontent.com/sakshamraj4/Abinbav_sustainability/main/fig.csv"
growth_data_df = pd.read_csv(growth_data_csv_url)
weather_data_url = "https://raw.githubusercontent.com/sakshamraj4/Abinbav_sustainability/main//weather_data.csv"
weather_df = pd.read_csv(weather_data_url)
activity_progress_url = 'https://github.com/sakshamraj4/Abinbav_sustainability/raw/main/crop_monitorig_protocol.csv'
activity_progress_df = pd.read_csv(activity_progress_url)
field_team_url = 'https://raw.githubusercontent.com/sakshamraj4/Abinbav_sustainability/main/field_team_data.csv'
field_team_df = pd.read_csv(field_team_url)
data_path = 'https://raw.githubusercontent.com/sakshamraj4/Abinbav_sustainability/main/risk_level.csv'
risk_summary_df = pd.read_csv(data_path)
GEOJSON_FILE_PATH = "https://raw.githubusercontent.com/sakshamraj4/Abinbav_sustainability/main/merged.geojson"

menu_options = ['Organisation level Summary', 'Plot level Summary', 'Map level View']
choice = st.sidebar.selectbox('Go to', menu_options)
if choice == 'Organisation level Summary':
    st.title("AB InBev Sustainability Dashboard")
    st.header("Organisation level Summarrization")    
    col1, col2, col3, col4, col5, col6 = st.columns(6)
    with col1:
        st.markdown('<div class="box"><h2>Total No of Plots</h2><p>59</p></div>', unsafe_allow_html=True)
    with col2:
        st.markdown('<div class="box"><h2>Total Area (Bigha)</h2><p>82.28</p></div>', unsafe_allow_html=True)
    with col3:
        st.markdown('<div class="box"><h2>Average Seed Rate</h2><p>8.01</p></div>', unsafe_allow_html=True)
    with col4:
        st.markdown('<div class="box"><h2>Average DAP/MOP Rate</h2><p>10.02</p></div>', unsafe_allow_html=True)
    with col5:
        st.markdown('<div class="box"><h2>Average Urea1 Rate</h2><p>7.01</p></div>', unsafe_allow_html=True)
    with col6:
        st.markdown('<div class="box"><h2>Average Urea2 Rate</h2><p>7.00</p></div>', unsafe_allow_html=True)   
    st.title("Crop Monitoring Observation")
    create_activity_progress_plot()   
    st.title('Risk Summary')
    plot_severity_counts(risk_summary_df, sort_by='specific_order')    
    st.header("Activity Progress")
    create_bar_plots(activity_df)    
    st.header("Growth Tracker Status")
    fig = create_stacked_bar_chart(growth_tracker_df)
    st.plotly_chart(fig)   
    st.download_button(
        label="Download Data updated via Field team",
        data=field_team_df.to_csv(index=False).encode('utf-8'),
        file_name='Daily_visit_data.csv',
        mime='text/csv'
    )   
    st.download_button(
        label="Download Data Entered via app",
        data=new_data_df.to_csv(index=False).encode('utf-8'),
        file_name='Field_team_updated_data.csv',
        mime='text/csv'
    )
    
elif choice == 'Plot level Summary':
    st.title("Plot level Summarization")
    st.header("Plot Visit Summary by Field Team")
    time_frame_options = ['Last Week', 'Last Month', 'Visit Till Date']
    selected_time_frame = st.selectbox('Select Time Frame', time_frame_options)   
    fig = create_line_chart(new_data_df, selected_time_frame)
    st.plotly_chart(fig)   
    st.download_button(
        label="Download Daily visit Data",
        data=new_data_df.to_csv(index=False).encode('utf-8'),
        file_name='Daily_visit_data.csv',
        mime='text/csv'
    )    
    st.subheader("Weather Trends")
    fig_weather = plot_weather_trends(weather_df)
    st.plotly_chart(fig_weather)
    
    st.download_button(
        label="Download Weather Data",
        data=weather_df.to_csv(index=False).encode('utf-8'),
        file_name='Weather_data.csv',
        mime='text/csv'
    )    
    st.header("Growth Tracker Data")
    fig = create_dot_plot(growth_tracker_df)
    st.plotly_chart(fig)   
    st.download_button(
        label="Download Growth Stage Data",
        data=growth_tracker_df.to_csv(index=False).encode('utf-8'),
        file_name='growth_tracker_data.csv',
        mime='text/csv'
    )   
    fig = create_dot_plot_1(growth_data_df)
    st.plotly_chart(fig)
    clicked_farm_name = st.experimental_get_query_params().get('farm_name', [None])[0]
    if clicked_farm_name:
        st.subheader(f"Detail view for Farm Name: {clicked_farm_name}")
        st.write(f"You clicked on {clicked_farm_name}. Add detailed view here.")      
    st.download_button(
        label="Download Activity Status Data",
        data=growth_data_df.to_csv(index=False).encode('utf-8'),
        file_name='activity_tracker_data.csv',
        mime='text/csv'
    )   
    st.title('Plot wise Risk Summary')
    fig = severity_dot_plot(risk_summary_df)
    st.plotly_chart(fig)
    
elif choice == 'Map level View':
    st.title("Map level View")
    st.sidebar.title("Options")
    zoom_level = st.sidebar.slider("Zoom Level", 1, 20, 15)
    # Change filter to Mature Stage
    mature_stage_filter = st.sidebar.multiselect("Filter by Mature Stage", ["Threshing", "Harvesting", "Not Harvesting"], default=["Threshing", "Harvesting", "Not Harvesting"])
    seed_variety_filter = st.sidebar.multiselect("Filter by Seed Variety", ["1207", "8214", "R&D Plot"], default=["1207", "8214", "R&D Plot"])

    st.markdown(
        """
        <style>
        body {
            font-family: 'Arial', sans-serif;
            background-color: #f9f9f9;
        }
        h1 {
            color: #333;
            margin-bottom: 20px;
        }
        .stButton>button {
            background-color: #007bff;
            color: white;
            border: none;
            padding: 12px 20px;
            text-align: center;
            text-decoration: none;
            display: inline-block;
            font-size: 16px;
            margin: 10px 0;
            cursor: pointer;
            border-radius: 5px;
            transition: background-color 0.3s ease;
        }
        .stButton>button:hover {
            background-color: #0056b3;
        }
        .download-button {
            background-color: #007bff;
            color: white;
            border: none;
            padding: 12px 20px;
            text-align: center;
            text-decoration: none;
            display: inline-block;
            font-size: 16px;
            margin: 10px 0;
            cursor: pointer;
            border-radius: 5px;
            transition: background-color 0.3s ease;
        }
        .download-button:hover {
            background-color: #0056b3;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    gdf = load_geojson(GEOJSON_FILE_PATH)
    if gdf is not None:
        # Update filtering to use Mature Stage
        if mature_stage_filter:
            gdf = gdf[gdf['Mature Stage'].isin(mature_stage_filter)]
        if seed_variety_filter:
            gdf = gdf[gdf['Seed-Varity'].isin(seed_variety_filter)]
        map_ = create_map(gdf, zoom_level)
        if map_:
            folium_static(map_, width=1200, height=800)
            zip_buffer = create_zip_file(gdf)
            zip_base64 = base64.b64encode(zip_buffer.getvalue()).decode()
            st.markdown(
                f"""
                <a class="download-button" href="data:application/zip;base64,{zip_base64}" download="all_plots_data.zip">Download All Data as ZIP</a>
                """,
                unsafe_allow_html=True
            )
